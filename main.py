import json
import re
import os
from datetime import datetime
from docx import Document
import questionary
from docx2pdf import convert

# Přidej styl pro questionary
from questionary import Style

custom_style = Style([
    ('qmark', 'fg:#ff9d00 bold'),
    ('question', 'fg:#aaaaaa bold'),  # změna na šedou
    ('answer', 'fg:#00ff00 bold'),
])

# Konstanty
FIRMY_PATH = 'firmy.json'
TEMPLATES_DIR = 'templates'
OUTPUT_DIR = 'Smlouvy'

# Načtení firem
def nacti_firmy():
    with open(FIRMY_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)

def vyber_firmu(firmy):
    vyber = questionary.select(
        "Vyber firmu:",
        choices=[f"{idx + 1}: {firma['nazev']}" for idx, firma in enumerate(firmy)]
    ).ask()
    idx = int(vyber.split(':')[0]) - 1
    return firmy[idx]

# Výběr šablony

def vyber_template():
    templates = [f for f in os.listdir(TEMPLATES_DIR) if f.endswith('.docx')]
    vyber = questionary.select(
        "Vyber typ smlouvy:",
        choices=[f"{idx + 1}: {template}" for idx, template in enumerate(templates)]
    ).ask()
    idx = int(vyber.split(':')[0]) - 1
    mesic = questionary.text("Zadej měsíc pro název souboru (např. 07):").ask()
    rok = questionary.text("Zadej rok pro název souboru (např. 2025):").ask()
    return os.path.join(TEMPLATES_DIR, templates[idx]), mesic, rok

def najdi_placeholdery(document):
    text = "\n".join([p.text for p in document.paragraphs])
    return set(re.findall(r"{(.*?)}", text))

def nahrad_placeholdery(document, hodnoty):
    for p in document.paragraphs:
        for key, value in hodnoty.items():
            if f"{{{key}}}" in p.text:
                p.text = p.text.replace(f"{{{key}}}", value)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in hodnoty.items():
                    if f"{{{key}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{key}}}", value)
    return document

def ziskej_hodnoty(placeholdery, firma):
    hodnoty = {}
    for placeholder in placeholdery:
        if placeholder.startswith("firma_"):
            key = placeholder.replace("firma_", "")
            hodnoty[placeholder] = firma.get(key, f"!!CHYBÍ {key}!!")
        else:
            print(f"\033[93m>>> {placeholder}\033[0m")
            odpoved = questionary.text(
                "Zadej hodnotu:",
                style=custom_style
            ).ask()
            hodnoty[placeholder] = odpoved
    return hodnoty

def uloz_smlouvu(document, hodnoty, mesic, rok):
    jmeno = hodnoty.get("jmeno", "Nezname")
    prijmeni = hodnoty.get("prijmeni", "Nezname")
    firma = hodnoty.get("firma_nazev", "Firma")
    nazev_smlouvy = hodnoty.get("nazev_smlouvy", "smlouva")
    
    filename = f"{prijmeni} {jmeno}_{firma}_{nazev_smlouvy}_{mesic}{rok}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    document.save(filepath)
    print(f"Smlouva uložena jako: {filepath}")
    if questionary.confirm("Chcete vytvořit i PDF verzi?").ask():
        pdf_filepath = filepath.replace('.docx', '.pdf')
        convert(filepath, pdf_filepath)
        print(f"PDF verze uložena jako: {pdf_filepath}")

def main():
    firmy = nacti_firmy()
    firma = vyber_firmu(firmy)
    template_path, mesic, rok = vyber_template()
    
    doc = Document(template_path)
    placeholdery = najdi_placeholdery(doc)
    
    hodnoty = ziskej_hodnoty(placeholdery, firma)
    
    nazev_smlouvy = os.path.splitext(os.path.basename(template_path))[0]
    hodnoty['nazev_smlouvy'] = nazev_smlouvy
    
    doc = nahrad_placeholdery(doc, hodnoty)
    uloz_smlouvu(doc, hodnoty, mesic, rok)

if __name__ == "__main__":
    main()